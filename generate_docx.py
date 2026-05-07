"""
REPORT.md → 한국 학교 보고서 양식 .docx 변환기 (서울시장 전용)
================================================================
- analyze.py 실행 결과(콘솔 출력)를 코드 블록으로 첨부
- charts/ 폴더의 PNG 이미지 4종을 보고서 본문에 임베드
- 한국 학교 보고서 양식 (한글에서도 열림)

실행: python generate_docx.py
출력: 도제학생 노희래_통계 분석 과제 보고서.docx
"""

import subprocess
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches

ROOT = Path(__file__).parent
CHARTS = ROOT / "charts"
OUTPUT = ROOT / "도제학생 노희래_통계 분석 과제 보고서.docx"
KO_FONT = "맑은 고딕"
MONO_FONT = "Consolas"


def set_korean_font(run, font=KO_FONT, size=10, bold=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font)


def add_paragraph(doc, text, font=KO_FONT, size=10, bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, color=None,
                  space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_korean_font(run, font=font, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {0: 24, 1: 16, 2: 13, 3: 11}
    bold = True
    color = RGBColor(0x1f, 0x3a, 0x6b) if level <= 2 else RGBColor(0x33, 0x33, 0x33)
    add_paragraph(doc, text, font=KO_FONT, size=sizes.get(level, 10), bold=bold,
                  color=color, space_before=14, space_after=8)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_korean_font(run, size=9, bold=True)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E7EEF7")
        tc_pr.append(shd)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_korean_font(run, size=9)
    return table


def add_console_block(doc, text):
    """analyze.py 콘솔 출력을 그대로 첨부 (모노스페이스, 회색 배경, 줄당 한 문단)"""
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(0.3)
        run = p.add_run(line if line else " ")
        run.font.name = MONO_FONT
        run.font.size = Pt(9)
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), MONO_FONT)
        rFonts.set(qn("w:hAnsi"), MONO_FONT)
        rFonts.set(qn("w:eastAsia"), KO_FONT)
        rPr.append(rFonts)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F4F6F8")
        pPr.append(shd)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = MONO_FONT
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), MONO_FONT)
    rFonts.set(qn("w:hAnsi"), MONO_FONT)
    rFonts.set(qn("w:eastAsia"), MONO_FONT)
    rPr.append(rFonts)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F8")
    pPr.append(shd)


def add_image(doc, image_path, caption, width_inches=6.2):
    """이미지 + 가운데 정렬 + 캡션"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    run = cap.add_run(caption)
    set_korean_font(run, size=9, color=RGBColor(0x55, 0x55, 0x55))


# =================== 보고서 생성 ===================

# analyze.py 실행 결과 캡처 (서울시장만)
print("  analyze.py 서울시장 실행 중...")
result = subprocess.run(
    ["python", str(ROOT / "analyze.py"), "서울시장"],
    capture_output=True, text=True, encoding="utf-8"
)
analyze_output = result.stdout

doc = Document()

# 마진
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# === 표지 ===
for _ in range(6):
    doc.add_paragraph()

add_paragraph(doc, "도제 과제 보고서", size=20, bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
add_paragraph(doc,
    "2026 서울시장 선거 당선 예측",
    size=26, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
    color=RGBColor(0x1f, 0x3a, 0x6b), space_after=8)
add_paragraph(doc,
    "여론조사 데이터를 활용한 통계적 추론",
    size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
    color=RGBColor(0x55, 0x55, 0x55), space_after=40)

for _ in range(8):
    doc.add_paragraph()

info_table = doc.add_table(rows=4, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ("학생명", "노희래"),
    ("과제명", "통계 분석 과제 (확률·통계 추론 통계학 적용)"),
    ("분석 도구", "Python 표준 라이브러리 (analyze.py)"),
    ("작성일", datetime.now().strftime("%Y년 %m월 %d일")),
]
for i, (label, val) in enumerate(info_data):
    cell0 = info_table.rows[i].cells[0]
    cell1 = info_table.rows[i].cells[1]
    cell0.text = ""; cell1.text = ""
    r0 = cell0.paragraphs[0].add_run(label)
    r1 = cell1.paragraphs[0].add_run(val)
    set_korean_font(r0, size=11, bold=True)
    set_korean_font(r1, size=11)
    cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_page_break()

# === 1. 과제 목표 ===
add_heading(doc, "1. 과제 목표와 접근 방식", 1)
add_paragraph(doc,
    "선관위·언론사가 수시로 발표하는 여론조사는 표본 일부(n=500~1000)로 "
    "모집단 전체(서울 유권자 약 800만 명)의 지지율을 추정하는 작업이다. "
    "PDF '머신러닝을 위한 통계 1' p8에서 본 추론 통계학의 정의 — "
    "\"수집된 데이터로 수집되지 않은 데이터를 설명\" — 그대로다.")
add_paragraph(doc,
    "이 과제에서 풀어야 할 핵심 질문은 두 가지였다.", space_before=4)
add_paragraph(doc,
    "  1) 단일 여론조사로 '몇 % 지지'라고 단정할 수 있는가? — 못한다. "
    "PDF p44 의 모비율 추정 공식대로 표본 비율 p 주변에 신뢰구간 "
    "p ± 1.96·√(pq/n) 가 있을 뿐이다.")
add_paragraph(doc,
    "  2) 여러 조사를 합치면 더 좁은 추정이 가능한가? — 가능하다. "
    "PDF p28 의 중심극한정리에 의해 σ 가 σ/√n 으로 줄어든다. "
    "표본을 4배로 합치면 신뢰구간 폭이 절반이 된다.")
add_paragraph(doc,
    "두 가지를 그대로 코드로 옮긴 게 analyze.py 다. 분석 대상은 "
    "전국적으로 가장 주목받는 서울시장 선거 (정원오 vs 오세훈) 다.",
    space_before=4)

# === 2. 통계 개념 이해 ===
add_heading(doc, "2. 학습한 통계 개념 — 내가 이해한 방식", 1)

add_heading(doc, "2.1 왜 1.96 인가 (PDF p15, p39, p47)", 2)
add_paragraph(doc, "처음에 1.96 이 어디서 오는지 헷갈렸다. PDF p15 표준정규분포표를 보고 정리하니 이렇다.")
add_paragraph(doc, "  · 표준정규분포 Z ~ N(0,1) 에서 P(0 ≤ Z ≤ 1.96) ≈ 0.475")
add_paragraph(doc, "  · 좌우 대칭이니 P(-1.96 ≤ Z ≤ 1.96) = 0.95")
add_paragraph(doc, "  · 즉 양쪽 꼬리(1종 오류 영역) 합 5%를 잘라낼 때 임계값이 |Z|=1.96")
add_paragraph(doc, "  · 99% 신뢰수준이면 이게 2.58 (PDF p47)")
add_paragraph(doc,
    "PDF p9 의 \"신뢰 = 욕먹지 않을 만큼\" 이라는 표현이 인상 깊었다. "
    "95% 신뢰구간은 \"100번 조사하면 95번은 이 구간 안에 진짜 모비율이 들어 있을 것이다\"이지, "
    "\"지금 이 한 조사가 95% 확률로 맞다\"가 아니다.", space_before=4)

add_heading(doc, "2.2 모비율 추정 — 선거 과제의 본진 (PDF p44)", 2)
add_paragraph(doc,
    "표기 약속: 본 보고서에서 'p' 는 표본에서 관측된 비율(PDF 표기로는 p-햇), "
    "'q' 는 1-p 를 의미한다. 모집단의 진짜 비율은 'p_모' 또는 문맥상 구별되는 "
    "표현을 사용한다.", space_before=4)
add_paragraph(doc, "PDF p44 는 정확히 이 과제용 슬라이드였다. 핵심 공식:")
add_code(doc, "  p - 1.96·√(pq/n)  ≤  p  ≤  p + 1.96·√(pq/n),  q = 1 - p")
add_paragraph(doc, "손으로 직접 계산해 봤다. 5/2 SBS 조사(입소스, n=800)에서 정원오 41% 였으니:")
add_code(doc,
    "  SE = √(0.41 × 0.59 / 800) = √0.000302 ≈ 0.01740\n"
    "  1.96 × 0.01740 ≈ 0.0341 = 3.41%p\n"
    "  CI(정원오) = [37.6%, 44.4%]")
add_paragraph(doc,
    "언론사가 발표한 오차범위 ±3.5%p 와 일치한다. 사실 언론사는 p=0.5 라는 "
    "최대 분산 가정으로 보수적으로 발표한다 (PDF p51 2010 수능 가형 30번 의 그 트릭이다 — "
    "p=1/2 일 때 신뢰구간이 가장 길어진다).")

add_heading(doc, "2.3 오차범위 내 접전 — Case A vs Case B (PDF p45)", 2)
add_paragraph(doc,
    "두 후보 비교에서 가장 중요한 개념. 두 후보 신뢰구간이 겹치면 표본 결과가 "
    "이 후보 우세더라도 모집단에서는 뒤집힐 수 있다 (Case B). 분리되면 "
    "표본 결과가 모집단에서도 그대로일 가능성이 매우 높다 (Case A).")
add_paragraph(doc, "  · 5/2 SBS 서울: 정원오 [37.6%, 44.4%] / 오세훈 [30.7%, 37.3%] → 분리, Case A")
add_paragraph(doc, "  · 12/15 MBC 서울: 정원오 [30.7%, 37.3%] / 오세훈 [32.7%, 39.3%] → 겹침, Case B 가능")

add_heading(doc, "2.4 두 후보 차이의 분산 — 다항분포", 2)
add_paragraph(doc,
    "여론조사는 정원오·오세훈·기타·부동층처럼 다항분포 4개 카테고리다. "
    "두 카테고리 비율의 차이를 보려면 다항분포 공분산을 반영해야 한다.")
add_code(doc, "  Var(p_A - p_B) = (p_A + p_B - (p_A - p_B)²) / n")
add_paragraph(doc,
    "Cov(p_A, p_B) = -p_A·p_B/n 이 음의 공분산이라는 게 직관적이다 — "
    "한 후보가 더 많이 나오면 다른 후보는 자동으로 적게 나오니까.")
add_paragraph(doc, "서울 5/2 SBS로 손계산:")
add_code(doc,
    "  Var(p_정 - p_오) = (0.41 + 0.34 - 0.07²) / 800\n"
    "                  = 0.7451 / 800 = 0.000931\n"
    "  SE_Δ = √0.000931 ≈ 0.0305 = 3.05%p\n"
    "  Z = 0.07 / 0.0305 ≈ 2.29")
add_paragraph(doc, "analyze.py 가 출력한 Z=+2.29 와 일치 ✓")

add_heading(doc, "2.5 가설 검정과 p-value (PDF p59~60)", 2)
add_paragraph(doc,
    "귀무가설 H₀: 두 후보 동률. 대립가설 H₁: 정원오 > 오세훈. "
    "검정통계량 Z = (p_정 - p_오) / SE_Δ 가 표준정규분포를 따른다고 가정하고, "
    "관측 Z 보다 더 극단적인 값이 나올 확률 1 - Φ(Z) 가 p-value 다.")
add_paragraph(doc,
    "Z=2.29 이면 p-value = 1 - Φ(2.29) ≈ 0.011. "
    "5/2 SBS 한 번의 조사만으로도 5% 유의수준에서 H₀ 기각 가능하다.", space_before=4)

add_heading(doc, "2.6 중심극한정리 — 합치면 좁아진다 (PDF p10, p28)", 2)
add_paragraph(doc,
    "이 과제의 결정적 트릭. 한 조사 n=800 의 SE 가 ±3.5%p 라면, "
    "4개 조사를 합쳐 n_eff=3200 이 되면 SE 가 ±1.7%p 로 정확히 절반이 된다 "
    "(σ/√n 이 σ/√(4n) = σ/(2√n)).")

# === 3. 데이터 ===
add_heading(doc, "3. 분석에 사용한 데이터", 1)
add_paragraph(doc, "서울시장 여론조사 5건 (2025-12-15 ~ 2026-05-02)")
add_table(doc,
    ["조사일", "의뢰처", "조사기관", "n", "정원오", "오세훈"],
    [
        ["2025-12-15", "MBC",       "코리아리서치", "800",  "34.0%", "36.0%"],
        ["2026-02-10", "MBC",       "코리아리서치", "800",  "40.0%", "36.0%"],
        ["2026-04-20", "펜앤마이크", "여론조사공정", "1000", "49.5%", "30.9%"],
        ["2026-04-28", "MBC",       "코리아리서치", "800",  "48.0%", "32.0%"],
        ["2026-05-02", "SBS",       "입소스",       "800",  "41.0%", "34.0%"],
    ])
add_paragraph(doc,
    "원자료는 polls.csv 에 정리. 통합 표본 n_eff = 4,200.", space_before=6)

# === 4. 분석 결과 (analyze.py 콘솔 출력 그대로) ===
add_heading(doc, "4. 분석 결과 — analyze.py 실행 결과", 1)
add_paragraph(doc,
    "다음은 PowerShell 에서 python analyze.py 서울시장 명령으로 실행한 콘솔 출력 그대로다.",
    space_after=8)
add_console_block(doc, analyze_output.strip())

# === 4-1 시계열 추세 차트 ===
add_heading(doc, "4.1 지지율 추세 + 95% 신뢰구간 밴드", 2)
add_paragraph(doc,
    "12월 시점 \"오차범위 내 박빙(Case B)\" 이던 구도가 4월 말부터 "
    "\"오차범위 밖 분리(Case A)\" 로 명확히 전환됐다. 시간이 흐를수록 정원오 우위가 강화되는 추세.")
add_image(doc, CHARTS / "trend.png",
          "그림 1. 서울시장 후보 지지율 추세 (메이저 4건, 95% CI 밴드)")

# === 4-2 CI 비교 차트 ===
add_heading(doc, "4.2 조사별 95% 신뢰구간 비교 (Forest Plot)", 2)
add_paragraph(doc,
    "두 에러바가 겹치면 박빙(PDF p45 Case B), 분리되면 우세 확정(Case A). "
    "12월·2월 조사는 두 에러바가 겹쳐 있고, 4월 말 이후 명확히 분리된다.")
add_image(doc, CHARTS / "ci_compare.png",
          "그림 2. 5건 조사별 95% 신뢰구간 비교")

# === 4-3 CLT 효과 ===
add_heading(doc, "4.3 중심극한정리 — Poll-of-polls 효과 (PDF p28)", 2)
add_paragraph(doc,
    "단일 조사 (n=800, MoE ±3.41%p) vs 합산 (n_eff=3,200, MoE ±1.70%p). "
    "표본을 합치면 σ/√n 효과로 분포 폭이 정확히 절반이 되어 두 후보 분포의 겹침이 거의 사라진다. "
    "이것이 PDF p28 중심극한정리의 시각적 의미다.")
add_image(doc, CHARTS / "clt_effect.png",
          "그림 3. 단일 조사 vs Poll-of-polls 분포 비교 — 표본을 합치면 분포 폭이 좁아진다")

# === 4-4 가설검정 시각화 ===
add_heading(doc, "4.4 가설검정 시각화 — p-value 의 의미 (PDF p59)", 2)
add_paragraph(doc,
    "표준정규분포 N(0,1) 위에서 관측 Z=+6.85 의 위치. 빨간 점선이 관측 위치이고, "
    "그 오른쪽 꼬리 면적이 p-value (≈ 3.73×10⁻¹²) 다. 이 값이 임계값 1.96 (α=0.05) 과 "
    "2.58 (α=0.01) 모두 한참 넘어선다. 1종 오류가 일어날 확률이 사실상 0 이라는 뜻으로, "
    "99% 신뢰수준에서도 H0(두 후보 동률) 을 기각할 충분한 증거다.")
add_image(doc, CHARTS / "hypothesis.png",
          "그림 4. 가설검정 — 관측 Z 와 p-value 영역 시각화")

# === 5. 핵심 통찰 ===
add_heading(doc, "5. 핵심 통찰 — 이 과제로 깨달은 것", 1)

add_heading(doc, "5.1 단일 조사는 신뢰하지 말고 합쳐라", 2)
add_paragraph(doc,
    "12/15 MBC 조사 한 건만 보면 오세훈이 36% vs 정원오 34% 로 우세였다. "
    "이 하나로 \"오세훈 우세\" 결론 내리면 통계 오용이다. "
    "이후 5개월간 다른 조사들이 누적되면서 통합 결과는 정원오 +9.19%p 우세로 명확해졌다. "
    "단일 표본의 변동성(PDF p28)을 모집단 진실로 착각하지 말 것 — 이게 첫 교훈이다.")

add_heading(doc, "5.2 표본수 vs 신뢰구간 폭", 2)
add_paragraph(doc,
    "단일 조사(n=800)의 SE가 ±3.41%p, 5건 합산(n=4,200)에서 ±0.74%p. "
    "표본수 5.25배에 SE는 √5.25=2.29배 좁아진다는 PDF p28 의 σ/√n 공식이 그대로 확인된다. "
    "데이터를 더 모으면 결론이 더 좁아진다.")

add_heading(doc, "5.3 시간 추세도 정보다", 2)
add_paragraph(doc,
    "정원오는 12월에 34%였다. 5월에는 41%다. 7%p 상승. "
    "각각의 SE 가 ±3.5%p 라 두 시점 차이가 ±5%p 보다 크면 진짜 변화로 봐야 한다. "
    "7%p 차이는 명확히 추세다.")

add_heading(doc, "5.4 의뢰처 편향이 보인다", 2)
add_paragraph(doc,
    "4/20 펜앤마이크(보수성향 매체) 조사는 정원오 49.5%, 오세훈 30.9%로 격차 18.6%p. "
    "4/28 MBC도 비슷한 시기에 16%p. 두 조사 모두 정원오 압승을 보인다. "
    "만약 펜앤마이크가 \"의뢰처 편향상 보수 후보가 더 잘 나오는 곳\" 인데도 정원오가 우세라면, "
    "이건 편향을 거스르는 강한 신호다.")

add_heading(doc, "5.5 PDF p9 \"욕먹지 않을 만큼\"의 의미", 2)
add_paragraph(doc,
    "처음에 이 표현이 너무 가벼워 보였는데, 분석을 마치고 다시 읽으니 정확하다. "
    "모집단 진짜 지지율을 단정할 수 없는 상황에서 우리가 할 수 있는 건 "
    "\"이 정도면 욕 안 먹을 만한 답\"을 내는 것이다. 95% 신뢰구간은 그 답이다.")

# === 6. 최종 예측 ===
add_heading(doc, "6. 최종 예측", 1)
add_table(doc,
    ["항목", "값"],
    [
        ["예측 당선자", "정원오 (더불어민주당)"],
        ["예상 득표율", "정원오 ≈ 42.8%, 오세훈 ≈ 33.6%"],
        ["95% 예측구간", "정원오 [41.3%, 44.3%], 오세훈 [32.2%, 35.1%]"],
        ["격차 / SE", "+9.19%p / 1.34%p"],
        ["검정통계량 Z", "+6.85"],
        ["p-value", "≈ 10⁻¹¹ 수준"],
        ["통계적 확신", "매우 높음 (99% 신뢰수준)"],
    ])

# === 7. 한계 ===
add_heading(doc, "7. 한계와 주의사항 — 솔직한 자기 평가", 1)

add_heading(doc, "7.1 한 표본으로 모집단 추정의 본질적 한계 (PDF p44)", 2)
add_paragraph(doc,
    "PDF p44가 명시했듯 \"샘플(표본)은 n개짜리 한 번만\". 5개 조사가 있어도 "
    "합치는 순간 결국 한 번의 큰 표본이다. 응답률이 10~12% 수준이라 "
    "비응답자의 정치성향이 응답자와 다를 경우 이 모든 추정이 편향된다.")

add_heading(doc, "7.2 양자대결만 다룸", 2)
add_paragraph(doc,
    "분석은 두 후보(민주당 vs 국민의힘)만 가정했다. 후보 등록(5월 중순) 이후 "
    "진보당·조국혁신당·무소속 등이 추가되면 다자 구도가 된다. "
    "양자대결에서 기타 후보로 빠질 표가 어느 쪽 기반인지에 따라 결과가 달라질 수 있다.")

add_heading(doc, "7.3 시간 가중치 미적용", 2)
add_paragraph(doc,
    "12월 조사와 5월 조사를 동등 가중. 5개월 전 민심이 지금과 같다는 건 비현실적이다. "
    "시간 감쇠(예: 30일 반감기)를 적용하면 결과가 약간 달라진다.")

add_heading(doc, "7.4 부동층 결집 시나리오 (PDF p60 2종 오류)", 2)
add_paragraph(doc,
    "부동층이 약 22% 수준이다. 이 부동층이 막판에 어느 쪽으로 결집하느냐에 따라 "
    "실제 결과는 예측과 달라질 수 있다. 부동층 22%가 7:3(오세훈:정원오)로 "
    "결집하면 오세훈이 +8.8%p 추가 확보해 정원오를 따라잡을 가능성도 있다.")
add_paragraph(doc,
    "p-value 가 매우 작다고 해서 \"결과가 무조건 이렇게 된다\"가 아니라 "
    "\"여론조사 시점의 응답자 분포가 전체 모집단의 분포와 다를 확률이 매우 작다\"라는 의미다. "
    "모집단의 마음이 6/3 까지 안 바뀐다는 보장은 다른 문제다.", space_before=4)

add_heading(doc, "7.5 공직선거법 §108 (공표금지)", 2)
add_paragraph(doc,
    "5/28 0시 ~ 6/3 18시 사이 새 여론조사 결과의 공표는 법으로 금지된다. "
    "본 보고서의 데이터는 5/2 가 마지막이다. 마지막 일주일의 변동은 본 분석에 반영되지 않는다.")

# === 8. 검증 계획 ===
add_heading(doc, "8. 6/3 이후 자체 검증 계획", 1)
add_paragraph(doc, "선거 결과가 발표되면 다음 항목으로 본 모형을 검증한다.")
add_paragraph(doc, "  □ 정원오 실제 득표율이 95% 신뢰구간 [41.3%, 44.3%] 안에 들어왔는가?")
add_paragraph(doc, "  □ 오세훈 실제 득표율이 95% 신뢰구간 [32.2%, 35.1%] 안에 들어왔는가?")
add_paragraph(doc, "  □ 정원오가 1위로 당선되었는가?")
add_paragraph(doc, "  □ 격차의 95% CI [+6.6%p, +11.8%p] 안에 실제 격차가 들어왔는가?")
add_paragraph(doc, "  □ 만약 빗나갔다면 그 원인은: 부동층 결집(2종 오류)? 비응답 편향? 의뢰처 편향?")
add_paragraph(doc,
    "정원오가 1위 맞히면 PDF p9 의 \"욕먹지 않을 만큼의 근거\" 로서 본 분석이 작동한 것이다. "
    "뒤집혔다면 §7의 한계 중 어느 것이 결정적이었는지가 후속 학습 주제다.",
    space_before=6)

# === 부록 ===
add_heading(doc, "부록 A. PDF 슬라이드와 본 보고서 매핑", 1)
add_table(doc,
    ["사용한 개념", "PDF 슬라이드", "본 보고서 위치"],
    [
        ["추론 통계학 정의",                       "p8",       "§1"],
        ["신뢰의 의미 (욕먹지 않을 만큼)",          "p9",       "§2.1, §5.5"],
        ["정규분포 / 중심극한정리",                 "p10, p28", "§2.6, §4.3"],
        ["표준정규분포·1.96·2.58",                "p15, p47", "§2.1"],
        ["모비율 신뢰구간 (이 과제 핵심)",         "p44",      "§2.2, §4"],
        ["오차범위 내 접전 / Case A·B",           "p45",      "§2.3, §4.2"],
        ["최대 허용 표본오차 (p=0.5)",           "p51",      "§2.2"],
        ["가설검정 / p-value",                    "p59",      "§2.5, §4.4"],
        ["1종·2종 오류",                          "p60",      "§2.5, §7.4"],
    ])

add_heading(doc, "부록 B. 손계산 검증 (5/2 SBS 서울 조사)", 1)
add_code(doc,
    "n = 800, p_정 = 0.41, p_오 = 0.34\n\n"
    "[1] 모비율 신뢰구간 (PDF p44)\n"
    "    SE_정 = √(0.41 × 0.59 / 800) = 0.01740\n"
    "    95% MoE = 1.96 × 0.01740 = 0.0341 = 3.41%p\n"
    "    95% CI(정) = [37.6%, 44.4%]   ← analyze.py 와 일치 ✓\n\n"
    "[2] 차이 분산 (다항분포)\n"
    "    Var(p_정 - p_오) = (0.41 + 0.34 - 0.07²) / 800\n"
    "                     = 0.7451 / 800\n"
    "                     = 0.0009314\n"
    "    SE_Δ = √0.0009314 = 0.0305 = 3.05%p\n\n"
    "[3] 가설 검정 (PDF p59)\n"
    "    Z = (0.41 - 0.34) / 0.0305 = 2.295\n"
    "    p-value = 1 - Φ(2.295) ≈ 0.0109\n\n"
    "[4] 결론\n"
    "    유의수준 5% 에서 H0 기각\n"
    "    → \"정원오가 오세훈보다 진짜 앞선다\" 로 판단할 통계적 근거 충분")

add_heading(doc, "부록 C. 분석 코드와 라이브 사이트", 1)
add_paragraph(doc, "  · 분석 코드: analyze.py (Python 표준 라이브러리만 사용)")
add_paragraph(doc, "  · 라이브 사이트: https://chldbswlsl.github.io/election/")
add_paragraph(doc, "  · 깃허브 저장소: https://github.com/chldbswlsl/election")
add_paragraph(doc, "  · 원자료: polls.csv (서울시장 5건 + 성남·부산 9건 = 총 14건)")

# 저장
try:
    doc.save(OUTPUT)
    print(f"  ✓ {OUTPUT.name} 생성 완료 ({OUTPUT.stat().st_size:,} bytes)")
except PermissionError:
    # 워드/한글에서 파일을 열어둔 경우 → 타임스탬프 붙여 새 이름으로 저장
    ts = datetime.now().strftime("%H%M%S")
    alt = OUTPUT.with_name(OUTPUT.stem + f"_v{ts}" + OUTPUT.suffix)
    doc.save(alt)
    print(f"  ⚠ 기존 파일이 열려있어 새 이름으로 저장: {alt.name}")
    print(f"     ({alt.stat().st_size:,} bytes)")
