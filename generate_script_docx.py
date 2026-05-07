"""
발표 스크립트 .docx 생성 — 슬라이드별 멘트, 시간 분배, 예상 질문
실행: python generate_script_docx.py
출력: 노희래_발표스크립트.docx (바탕화면 동시 복사)
"""

from datetime import datetime
from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).parent
OUT = ROOT / "노희래_발표스크립트.docx"
DESKTOP = Path("C:/Users/노희래/Desktop/노희래_발표스크립트.docx")
KO_FONT = "맑은 고딕"
MONO_FONT = "Consolas"


def set_korean_font(run, font=KO_FONT, size=10, bold=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color: run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font)


def add_paragraph(doc, text, size=10, bold=False, color=None,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=4, indent_cm=0):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    set_korean_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 11}
    color = RGBColor(0x1F, 0x3A, 0x6B) if level <= 2 else RGBColor(0x33, 0x33, 0x33)
    add_paragraph(doc, text, size=sizes.get(level, 10), bold=True,
                  color=color, space_before=14, space_after=8)


def add_slide_block(doc, num, title, time_sec, script_text):
    """슬라이드별 헤더 + 멘트 박스"""
    add_paragraph(doc, f"슬라이드 {num} — {title}",
                  size=12, bold=True, color=RGBColor(0x1F, 0x3A, 0x6B),
                  space_before=12, space_after=2)
    add_paragraph(doc, f"⏱  {time_sec}초",
                  size=10, color=RGBColor(0x6B, 0x72, 0x80),
                  space_after=4, indent_cm=0.5)
    # 멘트 박스
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F8")
    pPr.append(shd)
    run = p.add_run(script_text)
    set_korean_font(run, size=11)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_korean_font(run, size=10, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_korean_font(run, size=10)
    return table


# =================== 문서 생성 ===================
doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# 표지
add_paragraph(doc, "발표 스크립트", size=22, bold=True,
              color=RGBColor(0x1F, 0x3A, 0x6B),
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_paragraph(doc, "2026 서울시장 선거 당선 예측", size=18, bold=True,
              color=RGBColor(0x1F, 0x3A, 0x6B),
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_paragraph(doc, "통계적 추론 기반 당선 예측 발표용 멘트 모음",
              size=12, color=RGBColor(0x55, 0x55, 0x55),
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

# === 발표 개요 ===
add_heading(doc, "1. 발표 개요", 1)
add_paragraph(doc, "총 길이: 약 9분 + 질의응답 (10분 발표 권장)", bold=True)
add_paragraph(doc, "구성: 12 슬라이드 (PPTX 파일 별도)", space_after=10)

add_table(doc,
    ["#", "슬라이드 제목", "시간", "핵심 내용"],
    [
        ["1", "표지", "30초", "본인 소개, 발표 주제"],
        ["2", "한 줄 결론", "30초", "정원오 99% 신뢰수준 확신"],
        ["3", "과제 목표 — 두 핵심 질문", "1분", "단일 vs 통합 표본"],
        ["4", "통계 ① 모비율 신뢰구간", "1분30초", "공식 + 손계산 검증"],
        ["5", "통계 ② 가설검정 / p-value", "1분", "Z-검정, 단측 H1"],
        ["6", "분석 데이터 5건", "30초", "표본 합계 4,200명"],
        ["7", "결과 ① 추세 차트", "45초", "12월 박빙 → 5월 우세"],
        ["8", "결과 ② 신뢰구간 비교", "45초", "Forest Plot, Case A/B"],
        ["9", "결과 ③ 합치면 좁아진다", "45초", "중심극한정리 시각화"],
        ["10", "결과 ④ 가설검정 시각화", "45초", "Z=+6.85 위치"],
        ["11", "한계 5가지", "1분30초", "부동층 시나리오 강조"],
        ["12", "검증 + 마무리", "30초", "체크리스트 + URL 안내"],
    ])

# === 슬라이드별 스크립트 ===
add_heading(doc, "2. 슬라이드별 발표 스크립트", 1)

scripts = [
    (1, "표지", 30,
     "안녕하세요, 노희래입니다. 오늘 제가 발표할 주제는 2026년 6월 3일 지방선거 중 "
     "전국적으로 가장 주목받는 서울특별시장 선거의 당선 예측입니다.\n\n"
     "PDF 강의자료에서 배운 추론 통계 개념을 실제 발표된 여론조사 5건에 적용해서, "
     "본인 나름대로 6월 3일 결과를 사전에 예측해 봤습니다.\n\n"
     "결론부터 미리 말씀드리면 정원오 후보의 당선이 99% 신뢰수준에서 통계적으로 유의하게 예측됐습니다. "
     "어떻게 그 결과가 나왔는지 지금부터 차근차근 설명드리겠습니다."),

    (2, "한 줄 결론", 30,
     "이 슬라이드를 한 줄로 요약하면, 정원오 후보 당선이 매우 강한 통계적 근거로 예측된다는 것입니다.\n\n"
     "p-value 가 3.73 곱하기 10의 마이너스 12승, 즉 사실상 0 입니다. "
     "이 의미는 '두 후보가 동률이라는 가정 하에 이 정도 차이가 우연히 나올 확률이 사실상 0' 이라는 뜻입니다.\n\n"
     "예상 득표율은 정원오 42.83%, 오세훈 33.64%로 격차 +9.19%p 입니다. "
     "통합 표본은 4,200명으로, 단일 조사 표본 800명의 5.25배 규모입니다."),

    (3, "과제 목표", 60,
     "이 과제에서 제가 풀어야 할 핵심 질문은 두 가지였습니다.\n\n"
     "첫 번째 질문은, 한 번의 여론조사로 모집단 전체의 진짜 지지율을 단정할 수 있는가? "
     "정답은 못 한다는 것입니다. 표본 비율 p (피) 주변에는 항상 신뢰구간이 존재합니다. PDF p44 의 핵심 내용입니다.\n\n"
     "두 번째 질문은, 여러 조사를 합치면 더 정확해지는가? 정답은 가능하다는 것입니다. "
     "PDF p28 의 중심극한정리 덕분에, σ (시그마) 가 σ/√n (시그마 나누기 루트 엔) 으로 줄어들기 때문입니다. "
     "표본을 N배 늘리면 신뢰구간 폭은 √N (루트 엔) 배 좁아집니다.\n\n"
     "이 두 원리를 그대로 파이썬 코드, 즉 analyze.py 로 구현해서 서울시장 여론조사 5건을 분석했습니다."),

    (4, "통계 ① 모비율 신뢰구간", 90,
     "첫 번째 통계 개념은 모비율 신뢰구간입니다. "
     "공식 자체는 PDF p44 그대로입니다. "
     "p (피) 플러스마이너스 1.96 곱하기 루트, p (피) 곱하기 1 빼기 p, 나누기 n.\n\n"
     "처음에 1.96 이라는 숫자가 어디서 오는지 헷갈렸는데요, "
     "표준정규분포에서 양쪽 꼬리 합 5%를 자르는 임계값이 절댓값 Z (제트) 가 1.96 이라는 것을 PDF p15 에서 확인했습니다. "
     "그래서 95% 신뢰구간이 되는 거죠.\n\n"
     "PDF p9 에서 '신뢰는 욕먹지 않을 만큼' 이라는 표현이 인상 깊었습니다. "
     "95%는 100번 조사하면 95번 안에 진짜 모비율이 들어 있을 거라는, 사회가 합의한 기준선입니다.\n\n"
     "공식이 정말 작동하는지 손으로 직접 계산해 봤습니다. "
     "5월 2일 SBS 조사 데이터로 SE (에스이, 즉 표준오차) 계산하면 영점영일칠사, "
     "여기에 일점구육 곱하면 오차범위 플러스마이너스 삼점사일 퍼센트포인트. "
     "이게 언론사가 발표한 플러스마이너스 삼점오 퍼센트포인트와 일치합니다. "
     "신뢰구간도 코드 출력과 동일합니다. "
     "코드가 통계 이론대로 정확하게 작동한다는 검증입니다."),

    (5, "통계 ② 가설검정과 p-value", 60,
     "두 번째 통계 개념은 가설검정입니다. "
     "두 후보의 지지율이 통계적으로 정말 다른지 검증하는 절차입니다.\n\n"
     "귀무가설 H₀ (에이치 제로) 는 '두 후보 동률', "
     "대립가설 H₁ (에이치 원) 은 '정원오가 오세훈보다 우세' 로 설정했습니다. "
     "이건 단측 검정입니다. 데이터 보고 방향을 바꾸는 건 통계적으로 무효이기 때문에, "
     "분석 시작 전에 H₁ 을 고정했습니다.\n\n"
     "검정통계량 Z (제트) 는 두 후보 지지율 차이를 차이의 표준오차로 나눈 값입니다. "
     "여기서 차이의 분산은 다항분포 공분산을 반영해 "
     "p_A (피 에이) 더하기 p_B (피 비) 빼기, p_A 빼기 p_B 의 제곱, 나누기 n 으로 계산합니다.\n\n"
     "p-value (피 밸류) 는 1 빼기 Φ(Z) (파이 오브 제트), "
     "즉 표준정규분포의 누적확률입니다. "
     "PDF p60 정의 그대로 1종 오류가 일어날 확률을 의미합니다. "
     "차이가 없는데 있다고 잘못 발표할 위험을 의미하고요, 작을수록 H₀ 기각 근거가 강합니다. "
     "본 분석 결과는 사실상 0 입니다."),

    (6, "분석 데이터", 30,
     "사용한 데이터를 보여드리는 슬라이드입니다. "
     "작년 12월부터 올 5월까지 5건의 서울시장 여론조사를 모았습니다.\n\n"
     "MBC 의뢰 3건, SBS 1건, 펜앤마이크 1건으로 통합 표본수는 4,200명입니다. "
     "모두 중앙선거여론조사심의위에 등록된 정식 조사입니다.\n\n"
     "표 맨 아래 통합 행을 보시면 가중평균 지지율이 정원오 42.83%, 오세훈 33.64% 로 9.19%p 차이가 납니다. "
     "단일 조사들을 보면 시간이 갈수록 정원오의 우위가 굳혀지는 추세를 확인할 수 있습니다."),

    (7, "결과 ① 추세 차트", 45,
     "첫 번째 결과 차트는 시간에 따른 지지율 변화입니다. "
     "파란 선이 정원오, 빨간 선이 오세훈, 옅은 색 밴드가 95% 신뢰구간입니다.\n\n"
     "12월에는 두 밴드가 겹쳐서 PDF p45 의 Case B, 즉 박빙 상황이었습니다. "
     "이 시점에서는 통계적으로 우세를 단정할 수 없습니다.\n\n"
     "그러나 4월 말부터 두 신뢰구간이 완전히 분리됩니다. PDF p45 의 Case A, 우세 확정으로 전환된 거죠. "
     "추세상 정원오의 우위가 점점 굳혀지는 모습이 시각적으로 명확합니다."),

    (8, "결과 ② 신뢰구간 비교", 45,
     "두 번째 차트는 조사별 신뢰구간을 가로 막대로 보여주는 Forest Plot 입니다. "
     "각 행이 한 조사이고, 두 후보의 95% 신뢰구간이 서로 겹치는지 분리되는지가 핵심입니다.\n\n"
     "위쪽 두 행, 그러니까 2025년 12월과 2026년 2월 조사는 두 막대가 겹쳐 박빙입니다. "
     "아래 세 행은 두 막대가 분리되어 통계적 우세를 보입니다.\n\n"
     "오른쪽에 초록색 글씨로 '오차범위 밖' 이라고 적힌 게 우세가 확정된 조사들입니다. "
     "단일 조사 5건 중 3건이 이미 오차범위 밖이라는 점이 중요합니다."),

    (9, "결과 ③ 중심극한정리", 45,
     "세 번째 차트가 PDF p28 중심극한정리의 시각적 의미를 가장 잘 보여줍니다. "
     "왼쪽이 단일 조사 800명일 때의 분포, 오른쪽이 5건 합쳐 4,200명일 때의 분포입니다.\n\n"
     "왼쪽에서는 두 후보 분포가 살짝 겹치지만, 오른쪽에서는 완전히 분리됩니다. "
     "표본을 5.25배 늘리면 분포 폭이 √5.25 즉 2.29배 좁아져, 두 후보가 명확히 구분되는 거죠.\n\n"
     "이게 여러 조사를 통합 분석해야 하는 이유이자, 이번 분석의 결정적 트릭입니다. "
     "한 조사로는 못 보이던 차이가 합치면 명확해진다는 게 통계의 힘입니다."),

    (10, "결과 ④ 가설검정 시각화", 45,
     "마지막 차트는 가설검정의 시각적 표현입니다. "
     "표준정규분포 N (제로, 원), 즉 평균 0에 표준편차 1인 정규분포 위에서 "
     "관측 Z (제트) 가 어디 있는지 보여줍니다.\n\n"
     "Z (제트) 가 플러스 육점팔오 입니다. "
     "95% 신뢰수준 임계값 일점구육 과 99% 신뢰수준 임계값 이점오팔 을 모두 한참 넘었습니다. "
     "그래서 빨간 꼬리 면적, 즉 p-value (피 밸류) 가 사실상 0 인 겁니다.\n\n"
     "결론, H₀ (에이치 제로), 즉 '두 후보 동률' 가설을 99% 신뢰수준에서도 강하게 기각합니다. "
     "이게 본 분석의 통계적 근거입니다."),

    (11, "한계와 주의사항", 90,
     "여기까지가 결과였고, 이제 한계도 솔직하게 말씀드리겠습니다. "
     "5가지 잠재적 약점이 있습니다.\n\n"
     "첫째, 양자대결 가정. 후보 등록이 5월 중순이라 진보당이나 무소속 등이 추가되면 다자 구도가 됩니다.\n"
     "둘째, 비응답 편향. 응답률이 10에서 12% 수준이라 무응답자의 정치성향이 응답자와 다르면 전체가 편향됩니다.\n"
     "셋째, 시간 가중치 미적용. 12월 조사와 5월 조사를 동등 가중했는데, 사실 최근 조사가 더 의미 있습니다.\n"
     "넷째, House Effect 미보정. 의뢰처별 편향을 무시했습니다.\n\n"
     "그리고 노란 박스에 강조한 가장 큰 변수는 부동층 결집입니다. "
     "약 22% 부동층이 막판에 어느 쪽으로 흘러가느냐가 결과를 흔들 수 있습니다. "
     "비관 시나리오, 즉 부동층 7대3 으로 오세훈으로 결집하면 9%p 격차가 0.4%p 박빙까지 좁혀집니다.\n\n"
     "p-value 가 작다고 결과가 무조건 이렇게 된다는 건 아니라는 점, "
     "그리고 모집단의 마음이 6월 3일까지 안 바뀐다는 보장은 다른 문제라는 점을 인정합니다. "
     "통계는 '욕먹지 않을 만큼' 의 답이지 정답은 아니라는 PDF p9 의 의미를 분석 마치고 다시 깨달았습니다."),

    (12, "검증 + 마무리", 30,
     "마지막으로 6월 3일 이후 자체 검증 계획을 말씀드리겠습니다. "
     "체크리스트 4가지로, 두 후보 실제 득표율이 95% 신뢰구간 안에 들어왔는지, "
     "그리고 1위 적중인지가 핵심 검증입니다.\n\n"
     "만약 빗나간다면 한계 슬라이드의 어느 항목이 결정적이었는지가 후속 학습 주제가 될 것입니다. "
     "분석 코드와 인터랙티브 웹 사이트는 깃허브에 공개되어 있습니다.\n\n"
     "이상으로 발표를 마치겠습니다. 감사합니다. 질문 있으시면 답변드리겠습니다."),
]

for num, title, t, txt in scripts:
    add_slide_block(doc, num, title, t, txt)

# === 예상 질문 + 답변 ===
add_heading(doc, "3. 예상 질문 & 답변", 1)
qa = [
    ("왜 단측 검정을 했나요? 양측 검정이 더 일반적이지 않나요?",
     "단측 검정을 선택한 이유는 분석 시작 전 '민주당 후보 우세' 가설을 사전에 정해놓았기 때문입니다. "
     "양측 검정도 가능하지만, 여론조사 추세상 단측이 더 정보량 많은 검정이라 채택했습니다. "
     "데이터를 본 후 검정 방향을 바꾸는 건 통계적으로 무효(post-hoc analysis)라 처음부터 고정했습니다."),
    ("표본 5건이 충분한가요?",
     "각 조사 표본 크기가 800에서 1,000 사이로, 정규근사 조건 np > 5 와 n(1-p) > 5 를 모두 만족합니다. "
     "통합 표본은 4,200명이라 신뢰구간 폭이 ±1.5%p 수준으로 충분히 좁습니다. "
     "다만 시간 가중치를 적용 안 한 점은 한계로 인정합니다."),
    ("펜앤마이크 같은 편향성 매체 데이터를 왜 포함시켰나요?",
     "House Effect 보정을 안 한 게 분석의 한계지만, 보수성향 의뢰처도 정원오 우세를 보였다는 사실 자체가 "
     "'편향을 거스르는 강한 신호' 로 의미 있다고 판단해 포함시켰습니다. "
     "538.com 같은 곳은 의뢰처별 편향을 정량화해서 차감하는데, 그 부분은 후속 작업입니다."),
    ("p-value 3.73e-12 면 결과가 거의 확정 아닌가요? 왜 한계 슬라이드가 길어요?",
     "p-value 가 작다는 건 '여론조사 시점의 응답 분포가 동률 가정과 다를 확률이 작다' 는 의미이지, "
     "'6월 3일까지 모집단 마음이 안 바뀐다' 는 보장이 아닙니다. "
     "특히 부동층 22% 가 막판에 어느 쪽으로 결집하느냐는 통계로 못 잡습니다. "
     "그래서 비관 시나리오까지 함께 제시한 겁니다."),
    ("선거 결과가 빗나가면 분석이 잘못된 건가요?",
     "1위 적중인지가 핵심 검증입니다. 만약 빗나간다면 양자대결 가정·비응답 편향·부동층 결집 중 "
     "어느 항목이 결정적이었는지가 후속 학습 주제입니다. "
     "통계 모형은 '욕먹지 않을 만큼' 의 답이지 정답은 아니므로, 빗나가더라도 그 자체가 학습 자료입니다."),
    ("코드는 어떻게 만들었나요? 직접 작성하셨나요?",
     "Python 표준 라이브러리(csv, math) 만 사용했습니다. "
     "PDF p44 의 모비율 신뢰구간 공식, p59 의 가설검정 절차를 그대로 함수로 구현했습니다. "
     "손계산 결과와 코드 출력이 소수점 둘째 자리까지 일치한다는 걸 부록 B에서 검증했습니다. "
     "코드는 깃허브에 공개되어 있어서 누구나 검증 가능합니다."),
]
for q, a in qa:
    add_paragraph(doc, "Q. " + q, bold=True, color=RGBColor(0x1F, 0x3A, 0x6B),
                  size=11, space_before=10, space_after=2)
    add_paragraph(doc, "A. " + a, size=11, space_after=4, indent_cm=0.3)

# === 발표 시 팁 ===
add_heading(doc, "4. 발표 시 팁", 1)
tips = [
    "표지에서 '결론 미리 말씀드리면…' 으로 시작하면 청중 집중 ↑",
    "슬라이드 4 손계산 부분에서 칠판이나 종이에 직접 √(0.41×0.59/800) 풀어 보이면 임팩트 큼",
    "차트 슬라이드(7~10)에서는 '여기 파란 영역', '오른쪽 빨간 점선' 식으로 손가락·레이저 포인터로 가리키기",
    "한계 슬라이드(11)에서 '솔직히' 라는 표현 한 번 사용하면 신뢰감 ↑",
    "마지막에 '질문 있으시면…' 후 1~2초 침묵 가져가기 (자연스러운 끝맺음)",
    "PowerPoint 발표자 보기: F5 → Alt+F5 (현재 슬라이드 + 다음 슬라이드 + 노트 동시 표시)",
    "라이브 사이트 데모 가능하면 슬라이드 12 후 브라우저로 https://chldbswlsl.github.io/election/ 열기",
]
for t in tips:
    add_paragraph(doc, "• " + t, size=11, indent_cm=0.3, space_after=4)

# === 발음 가이드 ===
add_heading(doc, "5. 통계 기호 발음 가이드", 1)
add_paragraph(doc, "발표 중 자주 등장하는 기호·수식의 권장 발음.",
              size=11, color=RGBColor(0x55, 0x55, 0x55), space_after=8)

add_paragraph(doc, "5-1. 핵심 기호 — 본 발표에 직접 등장",
              size=12, bold=True, color=RGBColor(0x1F, 0x3A, 0x6B),
              space_before=6, space_after=4)
add_table(doc,
    ["기호", "권장 발음", "한국어 풀이"],
    [
        ["Z",       "제트",            "검정통계량"],
        ["H₀",      "에이치 제로",     "귀무가설"],
        ["H₁",      "에이치 원",       "대립가설"],
        ["p / p̂",  "피 / 피 햇",      "표본비율 (또는 모비율)"],
        ["q",       "큐",              "1 − p"],
        ["p_A, p_B", "피 에이, 피 비", "후보 A, B 의 비율"],
        ["p-value", "피 밸류 (또는 피 값)", "유의확률"],
        ["Φ(z)",    "파이 오브 제트",  "표준정규 누적함수 (대문자 Phi)"],
        ["σ",       "시그마",          "(모)표준편차"],
        ["σ/√n",    "시그마 나누기 루트 엔", "표본평균의 표준오차"],
        ["μ",       "뮤",              "모평균"],
        ["α",       "알파",            "유의수준 (보통 0.05)"],
        ["β",       "베타",            "2종 오류 확률"],
        ["Δ",       "델타",            "차이 (예: SE_Δ = 차이의 표준오차)"],
        ["N(0, 1)", "정규분포 제로 원",  "평균 0, 분산 1 인 정규분포"],
        ["n",       "엔",              "표본 수"],
        ["√",       "루트",            "제곱근"],
        ["√N",      "루트 엔",         ""],
        ["±",       "플러스마이너스",  ""],
        ["10⁻¹²",   "십의 마이너스 십이승",  ""],
        ["%p",      "퍼센트 포인트",   ""],
        ["≈",       "약 / 거의",       ""],
        ["|Z|",     "절댓값 제트",     "Z의 부호 무시"],
    ])

add_paragraph(doc, "5-1-bonus. 통계에 자주 나오는 그리스 문자 (참고)",
              size=12, bold=True, color=RGBColor(0x1F, 0x3A, 0x6B),
              space_before=10, space_after=4)
add_table(doc,
    ["대문자 / 소문자", "발음", "통계학 의미"],
    [
        ["Α / α",  "알파",     "유의수준 (1종 오류 허용 한계)"],
        ["Β / β",  "베타",     "회귀계수 / 2종 오류 확률"],
        ["Γ / γ",  "감마",     "감마분포"],
        ["Δ / δ",  "델타",     "차이 / 변화량"],
        ["Ε / ε",  "엡실론",   "오차항"],
        ["Ζ / ζ",  "제타",     "(자주 안 쓰임)"],
        ["Η / η",  "에타",     "효과 크기"],
        ["Θ / θ",  "세타",     "임의 모수"],
        ["Λ / λ",  "람다",     "포아송분포 평균률"],
        ["Μ / μ",  "뮤",       "모평균"],
        ["Π / π",  "파이",     "원주율 (3.14...)"],
        ["Ρ / ρ",  "로",       "상관계수"],
        ["Σ / σ",  "시그마",   "총합 (대문자) / 표준편차 (소문자)"],
        ["Τ / τ",  "타우",     "Kendall 의 τ"],
        ["Φ / φ",  "파이",     "표준정규 CDF (대문자)"],
        ["Χ / χ",  "카이",     "카이제곱 (χ²) 분포"],
        ["Ψ / ψ",  "프사이",   "(자주 안 쓰임)"],
        ["Ω / ω",  "오메가",   "표본공간 (Ω)"],
    ])
add_paragraph(doc,
    "★ 같은 글자도 대소문자에 따라 다르게 쓰임에 주의 — Σ(시그마)는 '총합', σ(시그마)는 '표준편차'.",
    size=10, color=RGBColor(0x6B, 0x72, 0x80), space_before=4, indent_cm=0.3)

add_paragraph(doc, "5-2. 약어 풀어서 말하기 (청중 친숙도 ↑)",
              size=12, bold=True, color=RGBColor(0x1F, 0x3A, 0x6B),
              space_before=10, space_after=4)
add_table(doc,
    ["약어", "발음 그대로", "권장 (한국어)"],
    [
        ["SE",  "에스이",   "표준오차"],
        ["MoE", "엠오이",   "오차범위"],
        ["CI",  "씨아이",   "신뢰구간"],
        ["CLT", "씨엘티",   "중심극한정리"],
    ])

add_paragraph(doc, "5-3. 수식 발음 예시",
              size=12, bold=True, color=RGBColor(0x1F, 0x3A, 0x6B),
              space_before=10, space_after=4)
examples = [
    ("p ± 1.96 · √(p(1-p)/n)",
     "피 플러스마이너스 일점구육 곱하기 루트, 피 곱하기 일 빼기 피, 나누기 엔"),
    ("Z = (p_정 − p_오) / SE_Δ",
     "제트는 피 정 빼기 피 오 나누기, 차이의 표준오차"),
    ("p-value = 1 − Φ(Z)",
     "피 밸류는 일 빼기 파이 오브 제트"),
    ("p-value = 3.73 × 10⁻¹²",
     "피 밸류 삼점칠삼 곱하기 십의 마이너스 십이승"),
    ("95% CI [37.6%, 44.4%]",
     "95% 신뢰구간 삼십칠점육 퍼센트에서 사십사점사 퍼센트"),
    ("Z = +6.85",
     "검정통계량 제트가 플러스 육점팔오"),
]
for orig, pron in examples:
    add_paragraph(doc, "• " + orig, size=11, bold=True, indent_cm=0.3,
                  space_after=2)
    add_paragraph(doc, "  → " + pron, size=11, indent_cm=0.6,
                  color=RGBColor(0x55, 0x55, 0x55), space_after=8)

add_paragraph(doc, "Tip: 약어가 처음 등장할 때만 한국어로 풀어주고, 이후엔 짧게 약어로 말해도 OK",
              size=10, color=RGBColor(0x6B, 0x72, 0x80),
              space_before=8, indent_cm=0.3)

# === 시간 압축안 ===
add_heading(doc, "6. 시간 압축안 (옵션)", 1)
add_paragraph(doc, "5분 발표 시 — 슬라이드 1, 2, 4, 7, 8, 11, 12 만 사용 (7장)",
              size=11, space_after=4)
add_paragraph(doc, "3분 발표 시 — 슬라이드 1, 2, 7~10 통합 1장, 11, 12 (5장)",
              size=11, space_after=4)
add_paragraph(doc, "10분 발표 — 12장 그대로 + 라이브 데모 1분", size=11)

try:
    doc.save(OUT)
    print(f"  ✓ {OUT.name} 생성 완료 ({OUT.stat().st_size:,} bytes)")
    target = OUT
except PermissionError:
    target = OUT.with_name(OUT.stem + f"_v{datetime.now().strftime('%H%M%S')}" + OUT.suffix)
    doc.save(target)
    print(f"  ⚠ 기존 파일 열림 — 대체 이름 저장: {target.name}")

try:
    shutil.copy(target, DESKTOP)
    print(f"  ✓ 데스크톱 복사: {DESKTOP}")
except PermissionError:
    print(f"  ⚠ 데스크톱 파일이 열려있어 복사 실패. 워드에서 닫고 다시 실행하세요.")
